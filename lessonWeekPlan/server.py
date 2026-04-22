import os
import io
import json
import base64
from flask import Flask, request, jsonify, send_file, send_from_directory
import anthropic
from google import genai as google_genai
from google.genai import types as genai_types
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__, static_folder='public', static_url_path='')

# ── AI 統一呼叫層 ─────────────────────────────────────────────────────────────

def _get_provider_and_key():
    """從 Header 取得 provider（gemini/anthropic）和 API Key。"""
    provider = request.headers.get('X-AI-Provider', 'gemini').lower()
    api_key  = request.headers.get('X-API-Key', '').strip()
    if not api_key:
        # 本機開發 fallback
        env_key = 'GEMINI_API_KEY' if provider == 'gemini' else 'ANTHROPIC_API_KEY'
        api_key = os.environ.get(env_key, '')
    if not api_key:
        raise ValueError('未提供 API Key，請在右上角設定')
    return provider, api_key


def call_ai(prompt, max_tokens=8192, image_b64=None, image_media_type='image/jpeg'):
    """統一 AI 呼叫介面，自動依 provider 使用 Gemini 或 Anthropic。"""
    provider, api_key = _get_provider_and_key()

    if provider == 'gemini':
        import time
        client = google_genai.Client(api_key=api_key)
        if image_b64:
            img_bytes = base64.b64decode(image_b64)
            img_part = genai_types.Part.from_bytes(data=img_bytes, mime_type=image_media_type)
            contents = [img_part, prompt]
        else:
            contents = prompt
        # 依序嘗試：主力模型 → 備援模型，遇到 quota 錯誤等候重試
        GEMINI_MODELS = ['gemini-2.0-flash', 'gemini-2.0-flash-lite']
        last_err = None
        for model_name in GEMINI_MODELS:
            for wait in [0, 15]:
                try:
                    if wait:
                        time.sleep(wait)
                    resp = client.models.generate_content(
                        model=model_name,
                        contents=contents
                    )
                    return resp.text
                except Exception as e:
                    last_err = e
                    err_str = str(e)
                    if '429' in err_str or 'RESOURCE_EXHAUSTED' in err_str:
                        continue   # quota → 等一下再試
                    if '404' in err_str or 'NOT_FOUND' in err_str:
                        break      # 這個模型不存在 → 試下一個
                    raise          # 其他錯誤直接往上拋
        raise last_err

    else:  # anthropic
        client = anthropic.Anthropic(api_key=api_key)
        if image_b64:
            content = [
                {'type': 'image', 'source': {
                    'type': 'base64', 'media_type': image_media_type, 'data': image_b64
                }},
                {'type': 'text', 'text': prompt}
            ]
        else:
            content = prompt
        msg = client.messages.create(
            model='claude-sonnet-4-6',
            max_tokens=max_tokens,
            messages=[{'role': 'user', 'content': content}]
        )
        return msg.content[0].text

# ── helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_cell_valign(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def _cell_para(cell, text, bold=False, size=9, center=False):
    """寫入儲存格文字，自動解析 {藍:...} 標記為藍色字。"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    import re
    parts = re.split(r'\{藍:(.*?)\}', text or '', flags=re.DOTALL)
    # parts: [plain, blue, plain, blue, ...]
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = '標楷體'
        if i % 2 == 1:  # 奇數 index = 藍色內容
            r.font.color.rgb = RGBColor(0x1A, 0x5F, 0xA8)
    return p

def _heading_para(doc, text, bold=True, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = '標楷體'
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    return p

def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(widths_cm[i] * 567)))  # 1cm ≈ 567 twips
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def _parse_ai_json(text):
    text = text.strip()
    if text.startswith('```'):
        lines = text.split('\n')
        text = '\n'.join(lines[1:])
    if text.endswith('```'):
        text = text.rsplit('```', 1)[0]
    return json.loads(text.strip())

# ── routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return send_from_directory('public', 'index.html')


@app.route('/api/analyze-image', methods=['POST'])
def analyze_image():
    try:
        data = request.json
        b64 = data.get('image', '')
        media_type = data.get('mediaType', 'image/jpeg')
        prompt = (
            '請仔細閱讀這張圖片，這是一份台灣國小的課程相關資料（可能是教材目錄、課綱說明、手寫筆記或課本封面等）。'
            '請將圖片中與「課程大綱、教學內容、單元名稱、教學目標」相關的文字內容整理成繁體中文的課程大綱描述，'
            '格式清晰、條列說明，方便後續 AI 生成週次課程計畫使用。'
            '若圖片中有單元名稱請列出；若有教學目標或重點也請摘要。只回覆整理後的課程大綱文字，不需要其他說明。'
        )
        text = call_ai(prompt, max_tokens=1024, image_b64=b64, image_media_type=media_type)
        return jsonify({'text': text.strip()})
    except ValueError as e:
        return jsonify({'error': str(e)}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500


STYLE_HINTS = {
    '繪本／故事書':   '在「引起動機」融入繪本或故事書，請具體說明書名（可為虛構示例）及閱讀後的討論問題。',
    '文學作品／小說': '在「自學／討論」引用文學作品段落，引發閱讀理解與批判思考討論。',
    '操作遊戲／實驗': '在「自學／討論」設計動手操作或小組遊戲活動，增加學習參與度。',
    '影片／歌謠':     '在「引起動機」或「統整」階段融入影片欣賞或歌謠吟唱，輔助概念建立。',
    '分組合作討論':   '「自學／討論」和「發表／統整」以小組合作方式進行，設計具體的討論任務。',
    '戶外教學／踏查': '適合的週次可在「自學／討論」安排戶外觀察或社區踏查活動。',
}

def _make_weekly_prompt(info, outline, start_week, end_week, total_weeks):
    """建立單批週次的 prompt（start_week ~ end_week），含 4 步驟教學活動。"""
    is_first = start_week == 1
    is_last  = end_week == total_weeks
    styles = info.get('teachingStyles', [])
    if styles:
        style_lines = '\n'.join(f'  • {STYLE_HINTS.get(s, s)}' for s in styles)
        teaching_style_note = f'\n教師選擇優先融入以下教學媒介，請在 4 步驟中具體體現：\n{style_lines}\n- 教師提問要開放性、能引發學生思考'
    else:
        teaching_style_note = '\n- 4 步驟活動可靈活運用繪本、故事、遊戲、討論等方式\n- 教師提問要開放性、能引發學生思考'
    notes = '1. 內容要循序漸進，符合學生發展\n2. 每週學習表現與學習內容請依十二年國教課綱撰寫\n3. 評量方式盡量多元\n4. 每個活動步驟請設計具體情境，教師提問要有引導思考的問句'
    if is_first:
        notes += '\n5. 第一週安排課程介紹或準備活動'
    if is_last:
        notes += f'\n6. 第 {total_weeks} 週安排總複習或成果發表'

    return f"""你是一位有豐富教學經驗的台灣國小教師，請根據以下資料設計第 {start_week} 週到第 {end_week} 週（共 {end_week - start_week + 1} 週）的詳細課程計畫。

基本資料：
- 學年度：{info.get('year', '115')} 學年度
- 年級：{info.get('grade', '')}年級
- 學期：第{info.get('semester', '1')}學期
- 科目：{info.get('subject', '')}
- 每週節數：{info.get('periodsPerWeek', '')}節
- 全學期共：{total_weeks}週
- 課程大綱：{outline}

每週 4 步驟的寫作要求（請務必做到）：
【1. 引起動機】
- 提供一個貼近學生生活的真實情境或故事、繪本、謎題
- 列出 1～2 個具體的佈題或問題情境（含數字或具體場景）
- 教師提問：寫出能引發好奇心的開放式問題（「……？」格式）

【2. 自學／討論】
- 說明學生的探究活動步驟（分組、操作、觀察、計算等）
- 列出 1 個具體的學習任務或遊戲規則
- 教師提問：寫出引導思考的追問（「……？」格式）

【3. 發表／統整】
- 說明學生分享方式（上台、黑板、小白板等）
- 教師歸納重點（寫出 2～3 個核心概念或口訣）
- 教師提問：寫出深化理解的問題（「……？」格式）

【4. 練習／評量】
- 設計一個有趣的任務名稱（如「挑戰王」、「闖關任務」）
- 說明具體練習內容（題型、數量、操作方式）
- 說明評量方式（觀察、口頭、紙筆、作品）

請以 JSON 格式回覆（不含說明文字，只回覆 JSON）：
{{
  "weeks": [
    {{
      "weekNum": {start_week},
      "unitTitle": "單元名稱（簡短標題）",
      "learningPerformance": "學習表現（依課綱填寫）",
      "learningContent": "學習內容（依課綱填寫）",
      "competencyIndicator": "對應素養指標代碼與說明",
      "step1_motivation": "【引起動機】\n情境：……\n佈題：……\n教師提問：「……？」",
      "step2_selfLearning": "【自學／討論】\n活動：……（步驟說明）\n任務／遊戲：……\n教師提問：「……？」",
      "step3_presentation": "【發表／統整】\n發表方式：……\n核心概念：① …… ② …… ③ ……\n教師提問：「……？」",
      "step4_practice": "【練習／評量】\n任務：《……》\n內容：……\n評量：……",
      "resources": "教學資源/學習策略",
      "assessment": "評量方式",
      "issues": "融入議題名稱（如：性別平等、品德教育；若無則填空字串）"
    }}
  ]
}}

教學風格提示：{teaching_style_note}

議題融入規則（重要）：
- 每 3～4 週至少安排一週融入議題（如性別平等、品德教育、環境教育、人權、生命教育、法治教育等）
- 有融入議題的週次，請在對應的活動步驟文字中，將與議題相關的句子用「{{藍:句子內容}}」包起來
- 例如：「活動：……{{藍:引導學生討論不同性別在運動中的平等參與機會}}……」
- issues 欄位填入議題名稱，steps 內容要真正體現議題精神，不只是附帶提及

注意：
{notes}"""


def _make_structure_prompt(info, outline, total_weeks):
    """生成課程架構表的 prompt。"""
    return f"""你是一位有豐富教學經驗的台灣國小教師。請根據以下資料，將全學期課程規劃為數個「單元」，以 JSON 格式回覆（只回覆 JSON，不含說明）：

基本資料：
- 科目：{info.get('subject', '')}，年級：{info.get('grade', '')}年級，學期：第{info.get('semester', '1')}學期
- 全學期共：{total_weeks}週
- 課程大綱：{outline}

格式：
{{
  "structure": [
    {{
      "unitNum": 1,
      "unitTitle": "單元名稱",
      "weeks": "1-5",
      "weekCount": 5,
      "description": "本單元主要內容概述（1-2句）",
      "competencyIndicators": "對應素養指標代碼，例如：數-E-A2、數-E-B1"
    }}
  ]
}}

注意：所有週次加總必須等於 {total_weeks} 週，單元數建議 4-6 個。"""


@app.route('/api/generate-weekly', methods=['POST'])
def generate_weekly():
    data = request.json
    info = data.get('basicInfo', {})
    outline = data.get('outline', '')
    total_weeks = int(info.get('weeks', 21))

    BATCH = 5
    all_weeks = []
    structure = []
    try:
        # 先生成課程架構
        struct_prompt = _make_structure_prompt(info, outline, total_weeks)
        struct_text = call_ai(struct_prompt, max_tokens=2048)
        struct_data = _parse_ai_json(struct_text)
        structure = struct_data.get('structure', [])

        # 分批生成週次活動
        for start in range(1, total_weeks + 1, BATCH):
            end = min(start + BATCH - 1, total_weeks)
            prompt = _make_weekly_prompt(info, outline, start, end, total_weeks)
            text = call_ai(prompt, max_tokens=6000)
            batch = _parse_ai_json(text)
            all_weeks.extend(batch.get('weeks', []))

        return jsonify({'weeks': all_weeks, 'structure': structure})
    except ValueError as e:
        return jsonify({'error': str(e)}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-goals', methods=['POST'])
def generate_goals():
    data = request.json
    info = data.get('basicInfo', {})
    weekly = data.get('weeklyPlan', [])
    outline = data.get('outline', '')

    sample = '\n'.join(f"第{w['weekNum']}週：{w.get('unitTheme','')}" for w in weekly[:6])

    # Subject code prefix for domain literacy
    subject_code_map = {
        '國語文': '國', '閩南語文': '閩', '客家語文': '客', '英語文': '英',
        '健康與體育': '健體', '生活課程': '生活', '社會': '社',
        '自然科學': '自', '藝術': '藝', '綜合活動': '綜', '台灣手語': '手語'
    }
    subj = info.get('subject', '')
    grade_group = 'E'  # Elementary school (國小)

    prompt = f"""你是一位有豐富教學經驗的台灣國小教師，請根據以下資料撰寫課程目標與核心素養。

基本資料：
- 科目：{subj}，年級：{info.get('grade','')}年級，學期：第{info.get('semester','1')}學期
- 課程大綱：{outline}

部分課程週次（前幾週範例）：
{sample}

請以 JSON 格式回覆（不含說明文字，只回覆 JSON）：
{{
  "goals": [
    "課程目標一（條列式，4至6項）",
    "課程目標二",
    "課程目標三",
    "課程目標四"
  ],
  "generalLiteracy": {{
    "A1": {{"checked": true, "content": "身心素質具體說明（若不相關設 checked:false，content 填空）"}},
    "A2": {{"checked": true, "content": ""}},
    "A3": {{"checked": false, "content": ""}},
    "B1": {{"checked": true, "content": ""}},
    "B2": {{"checked": false, "content": ""}},
    "B3": {{"checked": false, "content": ""}},
    "C1": {{"checked": true, "content": ""}},
    "C2": {{"checked": true, "content": ""}},
    "C3": {{"checked": false, "content": ""}}
  }},
  "domainLiteracy": "依{subj}領域綱要核心素養具體內涵填寫，格式：代碼-{grade_group}-XX 內涵說明（可多行）"
}}

總綱核心素養：A1=身心素質與自我精進, A2=系統思考與解決問題, A3=規劃執行與創新應變,
B1=符號運用與溝通表達, B2=科技資訊與媒體素養, B3=藝術涵養與美感素養,
C1=道德實踐與公民意識, C2=人際關係與團隊合作, C3=多元文化與國際理解"""

    try:
        text = call_ai(prompt, max_tokens=4096)
        result = _parse_ai_json(text)
        return jsonify(result)
    except ValueError as e:
        return jsonify({'error': str(e)}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    data = request.json
    info = data.get('basicInfo', {})
    weekly = data.get('weeklyPlan', [])
    goals = data.get('goals', [])
    gen_literacy = data.get('generalLiteracy', {})
    domain_literacy = data.get('domainLiteracy', '')
    structure = data.get('structure', [])

    doc = _build_docx(info, goals, gen_literacy, domain_literacy, weekly, structure)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{info.get('year','114')}學年度_{info.get('grade','')}年級_{info.get('subject','')}課程計畫.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ── docx builder ──────────────────────────────────────────────────────────────

def _build_docx(info, goals, gen_literacy, domain_literacy, weekly_plan, structure=None):
    doc = Document()

    # Page margins (A4, narrow)
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    # Default style
    doc.styles['Normal'].font.name = '標楷體'
    doc.styles['Normal'].font.size = Pt(10)

    year = info.get('year', '114')
    grade = info.get('grade', '')
    sem = info.get('semester', '1')
    subj = info.get('subject', '')
    designer = info.get('designer', '')
    ppw = info.get('periodsPerWeek', '')
    weeks = info.get('weeks', '21')
    try:
        total = int(ppw) * int(weeks)
    except Exception:
        total = ''
    mixed = info.get('mixedGrade', False)
    mixed_grades = info.get('mixedGrades', '')

    # ── 標題 ────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f'新北市海山國民小學{year}學年度　{grade}年級　第{sem}學期　部定課程計畫')
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = '標楷體'

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = d.add_run(f'設計者：{designer}')
    dr.font.size = Pt(11)
    dr.font.name = '標楷體'

    # ── 一、課程類別 ─────────────────────────────────────────────────────────
    _heading_para(doc, '一、課程類別：（請勾選）')

    subjects = [
        '國語文', '閩南語文', '客家語文', '英語文', '數學',
        '健康與體育', '生活課程', '社會', '自然科學',
        '藝術', '綜合活動', '台灣手語'
    ]
    line = '　'.join(
        f'{"☑" if subj == s else "□"} {s}' for s in subjects
    )
    p = doc.add_paragraph()
    r2 = p.add_run(line)
    r2.font.name = '標楷體'
    r2.font.size = Pt(10)

    # ── 二、學習節數 ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run(f'二、學習節數：每週（{ppw}）節，實施（{weeks}）週，共（{total}）節。')
    r.font.size = Pt(11)
    r.font.name = '標楷體'

    # ── 三、課程目標 ─────────────────────────────────────────────────────────
    _heading_para(doc, '三、課程目標：')
    for i, g in enumerate(goals, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f'（{i}）{g}')
        r.font.size = Pt(10)
        r.font.name = '標楷體'

    # ── 四、課程內涵 ─────────────────────────────────────────────────────────
    _heading_para(doc, '四、課程內涵：')

    # Sub-heading
    p = doc.add_paragraph()
    r = p.add_run('（一）總綱核心素養項目及具體內涵勾選')
    r.bold = True; r.font.size = Pt(10); r.font.name = '標楷體'
    p.paragraph_format.left_indent = Cm(0.5)

    literacy_labels = {
        'A1': 'A1 身心素質與自我精進',
        'A2': 'A2 系統思考與解決問題',
        'A3': 'A3 規劃執行與創新應變',
        'B1': 'B1 符號運用與溝通表達',
        'B2': 'B2 科技資訊與媒體素養',
        'B3': 'B3 藝術涵養與美感素養',
        'C1': 'C1 道德實踐與公民意識',
        'C2': 'C2 人際關係與團隊合作',
        'C3': 'C3 多元文化與國際理解',
    }

    # Literacy table: 3 columns
    lit_table = doc.add_table(rows=3, cols=3)
    lit_table.style = 'Table Grid'
    keys = list(literacy_labels.keys())
    for row_i in range(3):
        for col_i in range(3):
            idx = row_i * 3 + col_i
            key = keys[idx]
            cell = lit_table.rows[row_i].cells[col_i]
            info_item = gen_literacy.get(key, {})
            checked = info_item.get('checked', False)
            content = info_item.get('content', '')
            check_char = '☑' if checked else '□'
            label = literacy_labels[key]
            cell_text = f'{check_char} {label}'
            if checked and content:
                cell_text += f'\n{content}'
            _cell_para(cell, cell_text, size=9)
            _set_cell_valign(cell)

    p = doc.add_paragraph()
    r = p.add_run('（二）學習領域核心素養')
    r.bold = True; r.font.size = Pt(10); r.font.name = '標楷體'
    p.paragraph_format.left_indent = Cm(0.5)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r2 = p2.add_run(domain_literacy or '（請依各領域綱要核心素養具體內涵填寫）')
    r2.font.size = Pt(10)
    r2.font.name = '標楷體'

    # ── 五、課程架構 ─────────────────────────────────────────────────────────
    _heading_para(doc, '五、課程架構：')

    if structure:
        # 架構表：單元 | 單元名稱 | 週次 | 週數 | 內容概述 | 素養指標
        st = doc.add_table(rows=1 + len(structure), cols=6)
        st.style = 'Table Grid'
        HDR2 = 'D9E1F2'
        sh0 = st.rows[0].cells
        for ci, txt in enumerate(['單元', '單元名稱', '週次', '週數', '內容概述', '對應素養指標']):
            _cell_para(sh0[ci], txt, bold=True, size=9, center=True)
            _set_cell_bg(sh0[ci], HDR2)
        for ri, u in enumerate(structure):
            row = st.rows[ri + 1].cells
            _cell_para(row[0], f'單元{u.get("unitNum",ri+1)}', bold=True, size=9, center=True)
            _cell_para(row[1], u.get('unitTitle', ''), size=9)
            _cell_para(row[2], u.get('weeks', ''), size=9, center=True)
            _cell_para(row[3], str(u.get('weekCount', '')), size=9, center=True)
            _cell_para(row[4], u.get('description', ''), size=9)
            _cell_para(row[5], u.get('competencyIndicators', ''), size=9)
        _set_col_widths(st, [1.2, 2.5, 1.5, 1.0, 5.8, 3.0])
    else:
        doc.add_paragraph()

    # ── 六、混齡教學 ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run('六、本課程是否實施混齡教學：')
    r.bold = True; r.font.size = Pt(11); r.font.name = '標楷體'

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    txt = f'☑ 是（{mixed_grades}）　□ 否' if mixed else '□ 是（___年級和___年級）　☑ 否'
    r2 = p2.add_run(txt)
    r2.font.size = Pt(10); r2.font.name = '標楷體'

    # ── 七、素養導向教學規劃 ─────────────────────────────────────────────────
    _heading_para(doc, '七、素養導向教學規劃：')
    p_note = doc.add_paragraph()
    p_note.paragraph_format.left_indent = Cm(0.5)
    r_note = p_note.add_run('請以不同顏色標示：出版社(黑)、改編教材(紅)、議題融入(藍)、校本特色(自訂)、修改處(綠底黑字)')
    r_note.font.size = Pt(9); r_note.font.name = '標楷體'

    # 週次大表：週次 | 單元名稱 | 學習重點(2欄) | 素養指標 | 4步驟 | 資源 | 評量 | 融入議題
    # 共 12 欄
    HDR_BG = 'BDD7EE'
    num_rows = 2 + len(weekly_plan)
    tbl = doc.add_table(rows=num_rows, cols=12)
    tbl.style = 'Table Grid'

    HDRS_R0 = ['週次', '單元名稱', '學習重點', '', '素養指標',
               '1. 引起動機', '2. 自學／討論', '3. 發表／統整', '4. 練習／評量',
               '資源／策略', '評量方式', '融入議題']
    HDRS_R1 = ['', '', '學習表現', '學習內容', '', '', '', '', '', '', '', '']

    h0 = tbl.rows[0].cells
    h1 = tbl.rows[1].cells

    # Row 0 headers
    for ci, txt in enumerate(HDRS_R0):
        if txt:
            _cell_para(h0[ci], txt, bold=True, size=8, center=True)
        _set_cell_bg(h0[ci], HDR_BG)

    # Merge 學習重點跨 col 2-3
    h0[2].merge(h0[3])

    # Row 1 sub-headers
    for ci, txt in enumerate(HDRS_R1):
        if txt:
            _cell_para(h1[ci], txt, bold=True, size=8, center=True)
        _set_cell_bg(h1[ci], HDR_BG)

    week_nums = ['一','二','三','四','五','六','七','八','九','十',
                 '十一','十二','十三','十四','十五','十六','十七','十八','十九','二十','二十一']

    for i, w in enumerate(weekly_plan):
        row = tbl.rows[2 + i].cells
        wn = w.get('weekNum', i + 1)
        label = f'第{week_nums[wn-1] if wn <= len(week_nums) else wn}週'
        _cell_para(row[0],  label,                                   bold=True, size=8, center=True)
        _cell_para(row[1],  w.get('unitTitle', ''),                  size=8)
        _cell_para(row[2],  w.get('learningPerformance', ''),        size=8)
        _cell_para(row[3],  w.get('learningContent', ''),            size=8)
        _cell_para(row[4],  w.get('competencyIndicator', ''),        size=8)
        _cell_para(row[5],  w.get('step1_motivation', ''),           size=8)
        _cell_para(row[6],  w.get('step2_selfLearning', ''),         size=8)
        _cell_para(row[7],  w.get('step3_presentation', ''),         size=8)
        _cell_para(row[8],  w.get('step4_practice', ''),             size=8)
        _cell_para(row[9],  w.get('resources', ''),                  size=8)
        _cell_para(row[10], w.get('assessment', ''),                 size=8)
        _cell_para(row[11], w.get('issues', ''),                     size=8)
        _set_cell_valign(row[0])

    # 欄寬（總 18 cm）
    _set_col_widths(tbl, [1.0, 1.5, 1.5, 1.5, 1.5, 2.5, 2.5, 2.5, 2.5, 1.5, 1.5, 1.5])

    return doc


if __name__ == '__main__':
    print('課程計畫工具啟動中...')
    print('請開啟瀏覽器前往：http://localhost:5000')
    app.run(debug=False, port=5001, host='0.0.0.0')
